from docx import Document
from docx.shared import RGBColor
from Data import extract_text_from_pdf , clean_Response , terminal_image , clean_Ouput
from getCode import response
import time
import re
import io
from docx.shared import Pt

async def write_to_docx( file_bytes: bytes ,codeFontSize,questionFontSize,language:str) -> io.BytesIO:
    doc = Document()

    text = extract_text_from_pdf(file_bytes)

    questions = re.split(r'(?<!\d)(?=\d{1,2}\.\s)', text)
    questions = [q.strip() for q in questions if q.strip()]

    for index, line in enumerate(questions):
        start_time = time.time()

        heading = doc.add_heading(line, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(questionFontSize)

        APIRESPONSE = response(line, language).text
        print(f"APIRESPONSE = {APIRESPONSE}")
        cleanedResponse = clean_Response(APIRESPONSE,language)
        print(f"language : {language} " )
        # print(f"APIRESPONSE = {APIResponse}")
        print(f"cleanedREs = {cleanedResponse}")
        if cleanedResponse is not None:
            code = cleanedResponse[0]
            output = clean_Ouput(cleanedResponse[1])
        else:
            code = ""
            output = ""
        
        # Writting Code
        code_heading = doc.add_heading('Code:', level=3)
        for run in code_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(questionFontSize)
            
        code_para = doc.add_paragraph(code)
        for run in code_para.runs:
            run.font.size = Pt(codeFontSize)
        code_para.paragraph_format.space_after = 0
        code_para.paragraph_format.line_spacing = 1

        # Writting Output
        terminalImage = terminal_image(output)
        output_image = io.BytesIO()
        terminalImage.save(output_image, format='PNG')
        output_image.seek(0)
        Output_heading = doc.add_heading('Output:', level=3)
        for run in Output_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(questionFontSize)
        doc.add_picture(output_image)
        
        endtime = time.time()
        print(f"Time taken for question {index + 1}: {endtime - start_time:.2f} seconds")
        
        if index < len(questions) - 1:
            doc.add_page_break()

    # Save the doc to memory
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)

    return output_stream  # This is the in-memory .docx file

